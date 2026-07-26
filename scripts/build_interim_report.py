#!/usr/bin/env python3
"""Build the editable interim-report DOCX from the GitHub-readable Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = "0B2545"
BLUE = "2E74B5"
TEAL = "2A9D8F"
MUTED = "5C677D"
PALE = "F4F6F9"
GRID = "C8D3DD"
WHITE = "FFFFFF"
CONTENT_DXA = 9360


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError("table widths must sum to 9360 DXA")
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clean_math(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("E:V"):
        return "E:V → {0,1}^m"
    if stripped.startswith(r"L_{\max}(E)"):
        return "Lmax(E) = max_(u,v)∈ℰ dH(E(u),E(v))"  # noqa: RUF001
    if "substack" in text and "L^*" in text:
        return "L*free(G,m) = min over injective E:V→{0,1}^m of max_(u,v)∈ℰ dH(E(u),E(v))"  # noqa: RUF001
    if r"\frac{(a,b,c)}" in text:
        return "(a,b,c) / √(a²+b²+c²)"
    if "c_v" in text and r"\sum" in text:
        return "c_v = Σ_(j=0…m−1) 2^j b_(v,j)"  # noqa: RUF001
    if "Cartesian" in text and "L^*" in text:
        return "L*free = 2 < 3 = Lmax,Cartesian Gray"
    replacements = {
        "\\rightarrow": "→",
        "\\to": "→",
        "\\ge": "≥",
        "\\le": "≤",
        "\\in": "∈",
        "\\neq": "≠",
        "\\max": "max",
        "\\min": "min",
        "\\lceil": "⌈",
        "\\rceil": "⌉",
        "\\log_2": "log₂",
        "\\sqrt": "√",
        "\\sum": "Σ",
        "\\frac": "",
        "\\mathrm": "",
        "\\text": "",
        "\\substack": "",
        "\\mathcal E": "ℰ",  # noqa: RUF001
        "\\mathcal{E}": "ℰ",  # noqa: RUF001
        "\\{": "{",
        "\\}": "}",
        "\\,": " ",
        "\\ ": " ",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = text.replace(r"L^*_{\mathrm{free}}", "L*free")
    text = text.replace(r"L_{\max}", "Lmax")
    text = text.replace(r"d_H", "dH")
    text = text.replace(r"m_0", "m₀")
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = text.replace("^*", "*").replace("\\", "")
    text = text.replace("&", " ").replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def add_inline(paragraph, text: str) -> None:
    # Normalise common inline formula fragments before interpreting Markdown
    # emphasis. In particular, the asterisk in L^* is mathematical notation,
    # not the beginning of an italic span.
    text = text.replace(r"L^*_{\mathrm{free}}", "L*free")
    text = text.replace(r"L_{\max}", "Lmax")
    text = text.replace(r"d_H", "dH")
    pattern = re.compile(
        r"(\[([^\]]+)\]\((https?://[^)]+|\.\.?/[^)]+)\)"
        r"|\*\*([^*]+)\*\*|`([^`]+)`|\$([^$]+)\$|\\\((.+?)\\\)"
        r"|(?<![\w^])\*([^*\n]+)\*(?!\w))"
    )
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            plain = text[position:match.start()]
            if "\\" in plain:
                leading = " " if plain[:1].isspace() else ""
                trailing = " " if plain[-1:].isspace() else ""
                plain = leading + clean_math(plain) + trailing
            paragraph.add_run(plain)
        if match.group(2) is not None:
            label, target = match.group(2), match.group(3)
            if target.startswith("http"):
                add_hyperlink(paragraph, label, target)
            else:
                run = paragraph.add_run(label)
                set_run_font(run, color=BLUE)
        elif match.group(4) is not None:
            paragraph.add_run(match.group(4)).bold = True
        elif match.group(5) is not None:
            run = paragraph.add_run(match.group(5))
            set_run_font(run, name="Menlo", size=9.5, color=NAVY)
        elif match.group(6) is not None:
            run = paragraph.add_run(clean_math(match.group(6)))
            set_run_font(run, name="Cambria Math")
        elif match.group(7) is not None:
            run = paragraph.add_run(clean_math(match.group(7)))
            set_run_font(run, name="Cambria Math")
        elif match.group(8) is not None:
            paragraph.add_run(match.group(8)).italic = True
        position = match.end()
    if position < len(text):
        plain = text[position:]
        if "\\" in plain:
            leading = " " if plain[:1].isspace() else ""
            trailing = " " if plain[-1:].isspace() else ""
            plain = leading + clean_math(plain) + trailing
        paragraph.add_run(plain)


def new_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.extend((tabs, ind))
    level.extend((start, num_fmt, level_text, suffix, p_pr))
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_decimal_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num))
    p_pr.append(num_pr)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    if "Equation" not in [style.name for style in doc.styles]:
        equation = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        equation = doc.styles["Equation"]
    equation.font.name = "Cambria Math"
    equation.font.size = Pt(11)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(10)


def set_header_footer(section) -> None:
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SPHERE ENCODING  |  INTERIM RESEARCH REPORT")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = footer.add_run("Page ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(footer)


def add_cover(doc: Document) -> None:
    for _ in range(7):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("INTERIM RESEARCH REPORT")
    set_run_font(run, size=11, color=TEAL, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Beyond Cartesian Gray Codes")
    set_run_font(run, size=30, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Exact and Scalable Search for Locally Smooth\nBinary Encodings of the Sphere")
    set_run_font(run, size=16, color=BLUE)
    p.paragraph_format.space_after = Pt(36)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Accepted Stages 2-4 and current Stage 5 implementation status")
    set_run_font(run, size=11, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(90)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("26 July 2026")
    set_run_font(run, size=12, color=NAVY, bold=True)
    doc.add_page_break()

    doc.add_heading("Contents", level=1)
    entries = [
        "Abstract",
        "1. Motivation",
        "2. Formal problem",
        "3. Sphere discretisations",
        "4. Deterministic baselines",
        "5. Exact optimisation method",
        "6. Exact results",
        "7. Bounded larger instances",
        "8. Relationship to prior work",
        "9. Scalable search extension",
        "10. Interpretation and limitations",
        "11. Next steps",
        "Conclusion",
        "References and appendices",
    ]
    for entry in entries:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(entry)
    doc.add_page_break()


def table_widths(headers: list[str]) -> list[int]:
    column_count = len(headers)
    if "Neighbours" in headers:
        return [2100, 1450, 900, 900, 860, 3150]
    if column_count == 6:
        return [2100, 700, 2400, 900, 1400, 1860]
    patterns = {
        2: [3400, 5960],
        3: [3300, 3030, 3030],
        4: [3000, 1360, 2200, 2800],
        5: [2800, 900, 2200, 2060, 1400],
        6: [2400, 760, 2500, 1000, 1200, 1500],
    }
    if column_count in patterns:
        return patterns[column_count]
    base = CONTENT_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, block: list[str]) -> None:
    parsed = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in block]
    data = [parsed[0], *parsed[2:]]
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, table_widths(data[0]))
    repeat_table_header(table.rows[0])
    for row_index, values in enumerate(data):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, value)
            for run in p.runs:
                set_run_font(run, size=7.5 if len(data[0]) >= 5 else 8.5, color=WHITE if row_index == 0 else None, bold=row_index == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def set_picture_alt(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def build(markdown_path: Path, source_root: Path, output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    set_header_footer(doc.sections[0])
    add_cover(doc)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line.strip() == "## Abstract")
    lines = lines[start:]
    index = 0
    paragraph_buffer: list[str] = []
    active_numbering_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        p = doc.add_paragraph()
        if "`" in text:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, text)
        paragraph_buffer = []

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush_paragraph()
            active_numbering_id = None
            index += 1
            continue
        if line.startswith("## "):
            flush_paragraph()
            heading = line[3:]
            level = 1
            doc.add_heading(heading, level=level)
            index += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_heading(line[4:], level=2)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[([^]]+)\]\(([^)]+)\)", line)
        if image_match:
            flush_paragraph()
            alt, target = image_match.groups()
            path = (markdown_path.parent / target).resolve()
            picture = doc.add_picture(str(path), width=Inches(6.25))
            set_picture_alt(picture, alt)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
            index += 1
            continue
        if line.startswith("*Figure ") and line.endswith("*"):
            flush_paragraph()
            p = doc.add_paragraph(style="Caption")
            p.add_run(line[1:-1])
            index += 1
            continue
        if line.startswith("**Table "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            add_inline(p, line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip()):
            flush_paragraph()
            block = [line, lines[index + 1].strip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index].strip())
                index += 1
            add_markdown_table(doc, block)
            continue
        if line == "\\[":
            flush_paragraph()
            equation_lines = []
            index += 1
            while index < len(lines) and lines[index].strip() != "\\]":
                equation_lines.append(lines[index].strip())
                index += 1
            p = doc.add_paragraph(style="Equation")
            p.add_run(clean_math(" ".join(equation_lines)))
            index += 1
            continue
        if re.match(r"^- ", line):
            flush_paragraph()
            active_numbering_id = None
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            flush_paragraph()
            if active_numbering_id is None:
                active_numbering_id = new_decimal_numbering(doc)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            apply_decimal_numbering(p, active_numbering_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
            index += 1
            continue
        if line == "---":
            flush_paragraph()
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1
    flush_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Beyond Cartesian Gray Codes"
    doc.core_properties.subject = "Interim research report on locally smooth binary sphere encodings"
    doc.core_properties.author = "Alex Kolesnikov"
    doc.core_properties.keywords = "sphere, Gray code, Hamming distance, graph embedding, CP-SAT"
    doc.core_properties.comments = "Generated reproducibly from the version-controlled Markdown report."
    doc.save(output_path)
    print(f"wrote {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-root", type=Path, default=Path(__file__).resolve().parents[1])
    parser.add_argument("--markdown", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    source_root = args.source_root.resolve()
    markdown_path = (args.markdown or source_root / "report" / "Sphere_Encoding_Interim_Report.md").resolve()
    output_path = (args.output or source_root / "report" / "Sphere_Encoding_Interim_Report.docx").resolve()
    build(markdown_path, source_root, output_path)


if __name__ == "__main__":
    main()
